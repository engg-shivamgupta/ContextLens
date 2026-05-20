import io
import markdown
from pathlib import Path
from typing import Dict

from bs4 import BeautifulSoup
from docx import Document
from fastapi import UploadFile, HTTPException, status
from pypdf import PdfReader
import logging

import re

import logging

logger = logging.getLogger(__name__)

# Regex to find URLs that are NOT already in Markdown link format
# Negative lookbehind (?<!\]\() ensures we don't match (http...) part of existing [text](http...)
URL_PATTERN = re.compile(r'(?<!\]\()(https?://[^\s<>"]+|www\.[^\s<>"]+)')

class FileProcessingService:
    """A service dedicated to extracting text content from various file formats."""

    async def extract_text_from_file(self, file: UploadFile) -> Dict[str, str]:
        """
        Extracts text content from an uploaded file based on its extension.
        Returns a dictionary containing the title and content.
        """
        contents = await file.read()
        filename = file.filename
        file_ext = Path(filename).suffix.lower()

        try:
            if file_ext == ".pdf":
                text = self._extract_from_pdf(contents)
            elif file_ext == ".docx":
                text = self._extract_from_docx(contents)
            elif file_ext == ".html":
                text = self._extract_from_html(contents)
            elif file_ext == ".md":
                text = self._extract_from_md(contents)
            elif file_ext == ".txt":
                text = self._extract_from_txt(contents)
            else:
                raise HTTPException(
                    status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                    detail=f"Unsupported file type: {file_ext}",
                )
            
            # Use the filename (without extension) as the default title
            title = Path(filename).stem
            
            # Post-process: ensure all links are properly formatted for RAG
            text = self._post_process_text(text)

            return {"title": title, "content": text}

        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to process file: {filename}. Error: {str(e)}",
            )

    def _post_process_text(self, text: str) -> str:
        """
        Global clean-up and formatting for extracted text.
        Specifically ensures all plain URLs are converted to Markdown links.
        """
        if not text:
            return ""
        
        # Function to replace matched URL with markdown format
        def replace_link(match):
            url = match.group(0)
            # Ensure protocol for www links
            href = url if url.startswith('http') else f'https://{url}'
            return f'[{url}]({href})'

        # Apply regex to linkify plain URLs
        text = URL_PATTERN.sub(replace_link, text)
        return text

    def _resolve_pdf_object(self, obj):
        """Resolves indirect objects to their actual value."""
        if hasattr(obj, "get_object"):
            return obj.get_object()
        return obj

    def _extract_from_pdf(self, contents: bytes) -> str:
        """Extracts text from PDF file contents, including embedded links."""
        with io.BytesIO(contents) as pdf_file:
            reader = PdfReader(pdf_file)
            full_text = []
            
            for page in reader.pages:
                text = page.extract_text()
                
                # Extract links from annotations
                links = []
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        try:
                            annot_obj = self._resolve_pdf_object(annot)
                            
                            # Ensure it's a Link annotation
                            if annot_obj.get("/Subtype") == "/Link":
                                # Check for Action (URL)
                                if "/A" in annot_obj:
                                    action = self._resolve_pdf_object(annot_obj["/A"])
                                    if "/URI" in action:
                                        links.append(action["/URI"])
                        except Exception as e:
                            logger.warning(f"Failed to process annotation: {e}")
                            continue
                
                # Append links to the bottom of the page text if found
                if links:
                    # Filter out non-string links and deduplicate
                    valid_links = {link for link in links if isinstance(link, str)}
                    
                    # Regex fallback: Find links in plain text that might not have annotations
                    text_links = set(re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', text))
                    valid_links.update(text_links)
                    
                    if valid_links:
                        text += "\n\n**Links found on this page:**\n"
                        for link in valid_links:
                            text += f"- [{link}]({link})\n"
                
                full_text.append(text)
                
        return "\n".join(full_text)

    def _extract_from_docx(self, contents: bytes) -> str:
        """Extracts text from DOCX file contents, including embedded links and TABLES as Markdown."""
        from docx import Document as DocxDocument
        from docx.document import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        with io.BytesIO(contents) as docx_file:
            doc = DocxDocument(docx_file)
            full_text = []

            # Use iter_inner_content to process elements (Paragraphs and Tables) in order
            # Note: iter_inner_content() is not standard in all python-docx versions using Document object directly
            # We iterate through the body elements directly
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find the paragraph object corresponding to this element
                    # We have to search for it or wrap it
                    # Optimization: It's faster to just iterate paragraphs and tables if order wasn't critical
                    # But order IS critical.
                    
                    # Alternative safer approach: Iterate doc.iter_inner_content() if available, 
                    # but since it might not be, we'll try a simpler approach of iterating paragraphs and tables 
                    # based on their xml order.
                    pass
            
            # SIMPLER ROBUST APPROACH:
            # We will use the fact that doc.paragraphs and doc.tables are separate lists.
            # But we want combined order.
            # reliable way: iterate over doc.element.body and match with objects.
            
            def get_markdown_table(table):
                md_lines = []
                # extracting headers (assuming first row is header)
                if not table.rows: 
                    return ""
                    
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                md_lines.append("| " + " | ".join(headers) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    md_lines.append("| " + " | ".join(cells) + " |")
                    
                return "\n" + "\n".join(md_lines) + "\n"

            # Helper to extract text+links from a paragraph object
            def get_para_text(para):
                para_text = ""
                for child in para._element:
                    if child.tag.endswith('r'): # Run
                        if child.text: para_text += child.text
                    elif child.tag.endswith('hyperlink'): # Hyperlink
                        r_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if r_id:
                            try:
                                rel = doc.part.rels[r_id]
                                if rel.target_mode == 'External':
                                    url = rel.target_ref
                                    if url:
                                        display_text = ""
                                        for subchild in child:
                                            if subchild.tag.endswith('r') and subchild.text:
                                                display_text += subchild.text
                                        para_text += f" [{display_text}]({url}) "
                            except Exception: pass
                return para_text.strip()

            # Main iteration over document body elements
            for child in doc.element.body:
                if child.tag.endswith('p'):
                    # Create a Paragraph object from the element
                    para = Paragraph(child, doc)
                    text = get_para_text(para)
                    if text: full_text.append(text)
                
                elif child.tag.endswith('tbl'):
                    # Create a Table object from the element
                    table = Table(child, doc)
                    table_md = get_markdown_table(table)
                    if table_md: full_text.append(table_md)

        return "\n".join(full_text)

    def _extract_from_html(self, contents: bytes) -> str:
        """Extracts text from HTML file contents, preserving links as Markdown."""
        soup = BeautifulSoup(contents, "html.parser")
        
        # Convert tags to Markdown links: [text](href)
        for a in soup.find_all('a', href=True):
            markdown_link = f"[{a.get_text(strip=True)}]({a['href']})"
            a.replace_with(markdown_link)
            
        return soup.get_text(separator="\n", strip=True)

    def _extract_from_md(self, contents: bytes) -> str:
        """Extracts text from Markdown file contents by converting to HTML first."""
        html = markdown.markdown(contents.decode("utf-8"))
        return self._extract_from_html(html.encode("utf-8"))

    def _extract_from_txt(self, contents: bytes) -> str:
        """Extracts text from a plain text file."""
        return contents.decode("utf-8")


# Singleton instance
file_processing_service = FileProcessingService()